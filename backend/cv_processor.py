"""
CV Upload and Processing Module
Handles validation, storage, and text extraction from CV files
Adapted for Vercel Serverless environment
"""

import os
import mimetypes
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document

# Configuration
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), '.uploads')

# For Vercel serverless, use /tmp directory for temporary file storage
if os.environ.get('VERCEL'):
    UPLOAD_FOLDER = '/tmp'

# Ensure upload folder exists (might not be possible on Vercel, so use try/except)
try:
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
except Exception:
    pass


class CVUploadError(Exception):
    """Custom exception for CV upload errors"""
    def __init__(self, error_code, message, details=None):
        self.error_code = error_code
        self.message = message
        self.details = details or {}
        super().__init__(self.message)


def validate_file_extension(filename):
    """
    Validate that the file has an allowed extension
    
    Args:
        filename: The name of the file
        
    Returns:
        bool: True if valid, False otherwise
    """
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext in ALLOWED_EXTENSIONS


def validate_file_size(file):
    """
    Validate that the file size is within limits
    
    Args:
        file: The file object
        
    Returns:
        tuple: (is_valid, file_size_bytes)
        
    Raises:
        CVUploadError: If file exceeds size limit
    """
    # Get file size
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)  # Reset to beginning
    
    if file_size == 0:
        raise CVUploadError(
            'empty_file',
            'The uploaded file is empty',
            {'file_size_bytes': 0}
        )
    
    if file_size > MAX_FILE_SIZE_BYTES:
        raise CVUploadError(
            'file_too_large',
            f'File size exceeds maximum limit of {MAX_FILE_SIZE_MB}MB',
            {
                'max_size_mb': MAX_FILE_SIZE_MB,
                'file_size_mb': round(file_size / (1024 * 1024), 2)
            }
        )
    
    return True, file_size


def extract_text_from_pdf(file):
    """
    Extract text from PDF file
    
    Args:
        file: The file object
        
    Returns:
        str: Extracted text from PDF
        
    Raises:
        CVUploadError: If PDF is corrupted or unreadable
    """
    try:
        file.seek(0)
        pdf_reader = PdfReader(file)
        
        if not pdf_reader.pages:
            raise CVUploadError(
                'corrupted_file',
                'PDF file appears to be corrupted or empty',
                {'file_type': 'pdf'}
            )
        
        text = ""
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        
        if not text.strip():
            raise CVUploadError(
                'unreadable_pdf',
                'PDF file could not be processed. It may be image-based or corrupted.',
                {'file_type': 'pdf'}
            )
        
        return text
    except CVUploadError:
        raise
    except Exception as e:
        raise CVUploadError(
            'pdf_processing_error',
            'Failed to process PDF file',
            {'file_type': 'pdf', 'error_detail': str(e)}
        )


def extract_text_from_docx(file):
    """
    Extract text from DOCX file
    
    Args:
        file: The file object
        
    Returns:
        str: Extracted text from DOCX
        
    Raises:
        CVUploadError: If DOCX is corrupted or unreadable
    """
    try:
        file.seek(0)
        doc = Document(file)
        
        if not doc.paragraphs and not doc.tables:
            raise CVUploadError(
                'corrupted_file',
                'DOCX file appears to be empty or corrupted',
                {'file_type': 'docx'}
            )
        
        text = ""
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise CVUploadError(
                'empty_file',
                'DOCX file contains no readable text',
                {'file_type': 'docx'}
            )
        
        return text
    except CVUploadError:
        raise
    except Exception as e:
        raise CVUploadError(
            'docx_processing_error',
            'Failed to process DOCX file',
            {'file_type': 'docx', 'error_detail': str(e)}
        )


def extract_text_from_doc(file):
    """
    Attempt to extract text from legacy DOC file
    Note: DOC format is deprecated and support is limited
    
    Args:
        file: The file object
        
    Returns:
        str: Extracted text or error guidance
        
    Raises:
        CVUploadError: If DOC format is not supported
    """
    raise CVUploadError(
        'unsupported_format',
        'Legacy DOC format is not fully supported. Please use DOCX instead.',
        {'file_type': 'doc', 'suggestion': 'Convert to DOCX and try again'}
    )


def extract_text_from_file(file, filename):
    """
    Extract text from uploaded file based on file type
    
    Args:
        file: The file object
        filename: The filename
        
    Returns:
        str: Extracted text from the file
        
    Raises:
        CVUploadError: If file cannot be processed
    """
    file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    if file_ext == 'pdf':
        return extract_text_from_pdf(file)
    elif file_ext == 'docx':
        return extract_text_from_docx(file)
    elif file_ext == 'doc':
        return extract_text_from_doc(file)
    else:
        raise CVUploadError(
            'unsupported_format',
            f'File format ".{file_ext}" is not supported',
            {
                'uploaded_format': file_ext,
                'supported_formats': list(ALLOWED_EXTENSIONS)
            }
        )


def save_file(file, filename):
    """
    Save uploaded file to disk
    Note: On Vercel, files are saved to /tmp (ephemeral)
    
    Args:
        file: The file object
        filename: The filename
        
    Returns:
        str: Path to saved file
    """
    try:
        file.seek(0)
        filename = secure_filename(filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        return filepath
    except Exception as e:
        # On Vercel, file saving might fail due to ephemeral nature
        # This is not critical as we've already extracted the content
        return f"temporary-{secure_filename(filename)}"


def process_cv_upload(file, filename):
    """
    Main function to process CV upload
    Validates file and extracts text content
    
    Args:
        file: The file object from request
        filename: The filename
        
    Returns:
        dict: {
            'success': bool,
            'filename': str,
            'file_size_mb': float,
            'content': str (extracted text),
            'file_type': str
        }
        
    Raises:
        CVUploadError: If any validation or processing fails
    """
    # Validate extension
    if not validate_file_extension(filename):
        raise CVUploadError(
            'invalid_file_type',
            f'File type not supported. Please upload a PDF, DOCX, or DOC file.',
            {
                'uploaded_filename': filename,
                'supported_formats': list(ALLOWED_EXTENSIONS)
            }
        )
    
    # Validate file size
    is_valid, file_size = validate_file_size(file)
    
    # Get file extension
    file_ext = filename.rsplit('.', 1)[1].lower()
    
    # Extract text
    content = extract_text_from_file(file, filename)
    
    # Save file to uploads folder (best effort on Vercel)
    filepath = save_file(file, filename)
    
    return {
        'success': True,
        'filename': secure_filename(filename),
        'file_size_mb': round(file_size / (1024 * 1024), 2),
        'content': content[:5000],  # Return first 5000 chars as preview
        'file_type': file_ext,
        'full_path': filepath
    }
